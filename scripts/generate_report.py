#!/usr/bin/env python3
"""
日报生成脚本
生成 Markdown 和 Word 格式的日报
"""

import json
import re
from datetime import datetime, timedelta
from pathlib import Path
from typing import List, Dict

# 配置
OUTPUT_DIR = Path("~/DailyReports").expanduser()
DAYS_BACK = 7


def load_crawl_results(file_path: Path) -> List[Dict]:
    """加载抓取结果"""
    with open(file_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    return data.get("results", [])


def categorize_news(news_list: List[Dict]) -> Dict[str, List[Dict]]:
    """按类别分类新闻"""
    categories = {
        "国家部委": [],
        "省级政策": [],
        "市级政策": [],
        "区级政策": [],
        "行业协会": [],
        "项目申报": [],
        "其他": []
    }
    
    for news in news_list:
        category = news.get("category", "其他")
        
        if "国家" in category or "中华人民共和国" in category:
            categories["国家部委"].append(news)
        elif "省级" in category or "福建省" in category:
            categories["省级政策"].append(news)
        elif "市级" in category or "厦门市" in category:
            categories["市级政策"].append(news)
        elif "区级" in category or "思明" in category or "集美" in category:
            categories["区级政策"].append(news)
        elif "协会" in category or "行业" in category:
            categories["行业协会"].append(news)
        elif "申报" in category or "项目" in category:
            categories["项目申报"].append(news)
        else:
            categories["其他"].append(news)
    
    # 移除空类别
    return {k: v for k, v in categories.items() if v}


def generate_markdown(news_list: List[Dict], date_str: str) -> str:
    """生成 Markdown 格式的日报"""
    
    # 统计信息
    total = len(news_list)
    categorized = categorize_news(news_list)
    
    # 生成 Markdown
    md = f"""# 政府网站新闻政策日报

**报告日期**: {date_str}
**数据来源**: 40个政府网站
**生成时间**: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}

---

## 概览

| 统计项 | 数量 |
|--------|------|
| 监控网站数 | 40 |
| 新增新闻/政策 | {total} |
| 分类数量 | {len(categorized)} |

---

"""
    
    # 按类别输出
    for category, items in categorized.items():
        md += f"## {category} ({len(items)}条)\n\n"
        
        for item in items:
            md += f"### {item['title']}\n\n"
            md += f"- **来源**: {item['site_name']}\n"
            md += f"- **日期**: {item['date']}\n"
            md += f"- **链接**: {item['url']}\n"
            md += f"- **摘要**: {item['summary']}\n\n"
        
        md += "---\n\n"
    
    # 页脚
    md += f"""
---

**报告生成时间**: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
**下次生成时间**: 明天 06:45

---

*本报告由 OpenClaw Daily News Crawler 自动生成*
"""
    
    return md


def generate_word_doc(md_content: str, output_path: Path):
    """生成 Word 文档"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # 解析 Markdown 并转换为 Word
        lines = md_content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]
            
            # 标题
            if line.startswith('# '):
                heading = doc.add_heading(line[2:], level=0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=1)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=2)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=3)
            
            # 分隔线
            elif line == '---':
                doc.add_paragraph('_' * 50)
            
            # 表格
            elif line.startswith('| ') and '|' in line[2:]:
                # 简单处理表格，实际应用中需要更复杂的逻辑
                pass
            
            # 列表
            elif line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
            
            # 普通段落
            elif line.strip():
                doc.add_paragraph(line)
            
            i += 1
        
        # 保存文档
        doc.save(output_path)
        print(f"[INFO] Word 文档已生成: {output_path}")
        return True
        
    except ImportError:
        print("[WARN] python-docx 未安装，跳过 Word 文档生成")
        return False
    except Exception as e:
        print(f"[ERROR] Word 文档生成失败: {str(e)}")
        return False


def main():
    """主函数"""
    import sys
    
    # 获取输入文件
    if len(sys.argv) > 1:
        crawl_result_file = Path(sys.argv[1])
    else:
        # 查找最新的抓取结果文件
        result_files = sorted(Path(".").glob("crawl_result_*.json"), reverse=True)
        if result_files:
            crawl_result_file = result_files[0]
        else:
            print("[ERROR] 未找到抓取结果文件")
            sys.exit(1)
    
    # 获取日期
    date_str = datetime.now().strftime("%Y-%m-%d")
    
    print(f"[INFO] 加载抓取结果: {crawl_result_file}")
    news_list = load_crawl_results(crawl_result_file)
    print(f"[INFO] 共 {len(news_list)} 条新闻")
    
    # 生成 Markdown
    print("[INFO] 生成 Markdown 文档...")
    md_content = generate_markdown(news_list, date_str)
    md_file = OUTPUT_DIR / "md" / f"{date_str}-日报.md"
    with open(md_file, "w", encoding="utf-8") as f:
        f.write(md_content)
    print(f"[INFO] Markdown 已保存: {md_file}")
    
    # 生成 Word 文档
    print("[INFO] 生成 Word 文档...")
    doc_file = OUTPUT_DIR / "doc" / f"{date_str}-日报.docx"
    success = generate_word_doc(md_content, doc_file)
    if success:
        print(f"[INFO] Word 文档已保存: {doc_file}")
    
    # 输出摘要
    categorized = categorize_news(news_list)
    print("\n" + "=" * 60)
    print("[INFO] 日报生成完成")
    print(f"[INFO] 总新闻数: {len(news_list)}")
    print(f"[INFO] 分类数量: {len(categorized)}")
    for cat, items in categorized.items():
        print(f"  - {cat}: {len(items)}条")
    print("=" * 60)


if __name__ == "__main__":
    main()
